import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime

st.set_page_config(
    page_title="HAS Manual Generator",
    page_icon="🌙",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600&family=DM+Serif+Display&display=swap');
    
    html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
    
    .main { background: #F8F6F0; }
    .block-container { max-width: 860px; padding-top: 2rem; padding-bottom: 3rem; }
    
    h1, h2, h3 { font-family: 'DM Serif Display', serif; }
    
    .step-header {
        background: white;
        border: 1px solid #E8E4DC;
        border-radius: 12px;
        padding: 1.25rem 1.5rem;
        margin-bottom: 1.5rem;
    }
    .step-num {
        font-size: 11px;
        color: #888780;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin-bottom: 4px;
    }
    .step-title {
        font-family: 'DM Serif Display', serif;
        font-size: 22px;
        color: #2C2C2A;
    }
    
    .info-box {
        background: #E1F5EE;
        border-left: 3px solid #1D9E75;
        border-radius: 0 8px 8px 0;
        padding: 10px 14px;
        font-size: 13px;
        color: #085041;
        margin-bottom: 1rem;
    }
    
    .warning-box {
        background: #FAEEDA;
        border-left: 3px solid #BA7517;
        border-radius: 0 8px 8px 0;
        padding: 10px 14px;
        font-size: 13px;
        color: #633806;
        margin-bottom: 1rem;
    }
    
    .summary-card {
        background: white;
        border: 1px solid #E8E4DC;
        border-radius: 12px;
        padding: 1.25rem 1.5rem;
        margin-bottom: 1rem;
    }
    .summary-label {
        font-size: 11px;
        color: #888780;
        text-transform: uppercase;
        letter-spacing: 0.06em;
        margin-bottom: 2px;
    }
    .summary-value {
        font-size: 15px;
        color: #2C2C2A;
        font-weight: 500;
    }
    
    .badge-has {
        display: inline-block;
        background: #FAEEDA;
        color: #633806;
        font-size: 12px;
        padding: 3px 10px;
        border-radius: 20px;
        font-weight: 500;
    }
    .badge-ihcs {
        display: inline-block;
        background: #E6F1FB;
        color: #0C447C;
        font-size: 12px;
        padding: 3px 10px;
        border-radius: 20px;
        font-weight: 500;
    }
    
    .progress-container {
        display: flex;
        gap: 6px;
        margin-bottom: 2rem;
    }
    
    div[data-testid="stButton"] > button {
        border-radius: 8px;
        font-weight: 500;
        font-size: 14px;
    }
    
    div[data-testid="stSelectbox"] > div { border-radius: 8px; }
    div[data-testid="stTextInput"] > div > div { border-radius: 8px; }
    div[data-testid="stTextArea"] > div { border-radius: 8px; }
    div[data-testid="stMultiSelect"] > div { border-radius: 8px; }
    
    footer { display: none; }
</style>
""", unsafe_allow_html=True)

# ─── Data ─────────────────────────────────────────────────────────────────────

INDUSTRIES = {
    "Makanan & Minuman": {"ms": "MS 2200", "code": "food"},
    "Kosmetik": {"ms": "MS 2634", "code": "cosmetic"},
    "Farmaseutikal": {"ms": "MS 2424", "code": "pharma"},
    "Logistik / Penyimpanan / Pengangkutan": {"ms": "MS 2400", "code": "logistics"},
    "Rumah Sembelih": {"ms": "MS 1500", "code": "slaughter"},
    "Katering / Restoran": {"ms": "MyTrust Catering", "code": "catering"},
    "Barang Gunaan": {"ms": "MS 2565", "code": "consumer"},
    "Peranti Perubatan": {"ms": "MS 2424-2", "code": "medical"},
}

COMPANY_SIZES = {
    "Mikro / Kecil (pekerja < 75, jualan tahunan < RM15j)": "micro",
    "Sederhana / SME (pekerja 76–200, jualan RM15j–RM50j)": "sme",
    "Besar / Multinasional (pekerja > 200 atau jualan > RM50j)": "large",
}

MANDATORY_SOPS = [
    ("audit", "SOP Audit Halal Dalaman"),
    ("risk", "SOP Kawalan Risiko Halal (Halal Control Point / HCP)"),
    ("rawmat", "SOP Kawalan Bahan Mentah"),
    ("training", "SOP Latihan Halal"),
    ("trace", "SOP Kebolehkesanan (Traceability)"),
    ("review", "SOP Semakan HAS"),
    ("lab", "SOP Analisis Makmal"),
    ("sertu", "SOP Sertu"),
]

RECORDS = {
    "audit": ["Rekod Audit Halal Dalaman (bulanan)", "Laporan Ketidakakuran (NCR)", "Pelan Tindakan Pembetulan"],
    "risk": ["Halal Risk Management Plan", "Rekod Pemantauan HCP (harian)", "Rekod Tindakan Pembetulan"],
    "rawmat": ["Raw Material Masterlist", "Rekod Kawalan Bahan Mentah (harian)", "Log Penerimaan Bahan Mentah"],
    "training": ["Rekod Kehadiran Latihan", "Rekod Penilaian Latihan", "Pelan Latihan Tahunan"],
    "trace": ["Rekod Kebolehkesanan Produk", "Prosedur Penarikan Produk (Recall)", "Diagram Kebolehkesanan"],
    "review": ["Minit Mesyuarat Semakan HAS (tahunan)", "Laporan Status HAS"],
    "lab": ["Rekod Analisis Makmal", "Sijil CoA Bahan Mentah"],
    "sertu": ["Rekod Pelaksanaan Sertu", "Prosedur Sertu"],
}

# ─── Session State Init ────────────────────────────────────────────────────────

def init_state():
    defaults = {
        "step": 1,
        "company_name": "",
        "company_address": "",
        "company_ssm": "",
        "company_size": None,
        "industry": None,
        "product_desc": "",
        "ihc_members": [],
        "halal_exec": "",
        "halal_policy": "",
        "selected_sops": [],
        "hcp_processes": [],
        "doc_controller": "",
        "approval_name": "",
        "approval_title": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ─── Helpers ──────────────────────────────────────────────────────────────────

def get_system_type():
    size_code = COMPANY_SIZES.get(st.session_state.company_size, "")
    return "IHCS" if size_code == "micro" else "HAS"

def get_ms_standard():
    ind = st.session_state.industry
    if ind and ind in INDUSTRIES:
        return INDUSTRIES[ind]["ms"]
    return "MS 2200"

def progress_bar():
    steps = ["Profil", "Industri", "JHD", "SOP & HCP", "Jana"]
    current = st.session_state.step
    cols = st.columns(len(steps))
    for i, (col, label) in enumerate(zip(cols, steps), 1):
        with col:
            if i < current:
                st.markdown(f'<div style="text-align:center; font-size:12px; color:#1D9E75; font-weight:500; padding:6px 0; border-bottom: 2px solid #1D9E75;">{label}</div>', unsafe_allow_html=True)
            elif i == current:
                st.markdown(f'<div style="text-align:center; font-size:12px; color:#2C2C2A; font-weight:600; padding:6px 0; border-bottom: 2px solid #2C2C2A;">{label}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div style="text-align:center; font-size:12px; color:#B4B2A9; padding:6px 0; border-bottom: 1px solid #E8E4DC;">{label}</div>', unsafe_allow_html=True)

def header(title, step_num, total=5):
    st.markdown(f"""
    <div class="step-header">
        <div class="step-num">Langkah {step_num} daripada {total}</div>
        <div class="step-title">{title}</div>
    </div>
    """, unsafe_allow_html=True)

# ─── Word Document Generator ──────────────────────────────────────────────────

def generate_has_manual():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    def set_heading(para, text, level=1, color=None):
        para.clear()
        run = para.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)
        if color:
            run.font.color.rgb = RGBColor(*color)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def add_section_header(doc, text, bg_color=(31, 158, 117)):
        para = doc.add_paragraph()
        run = para.add_run(f"  {text}  ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*bg_color))
        para._p.get_or_add_pPr().append(shd)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        return para

    def add_row(table, label, value):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    s = st.session_state
    system = get_system_type()
    ms = get_ms_standard()
    today = datetime.now().strftime("%d %B %Y")

    # ── Cover Page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"MANUAL {system}")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(15, 110, 86)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("SISTEM JAMINAN HALAL" if system == "HAS" else "SISTEM KAWALAN HALAL DALAMAN")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = company_para.add_run(s.company_name.upper() or "NAMA SYARIKAT")
    cr.bold = True
    cr.font.size = Pt(20)

    doc.add_paragraph()
    ref_table = doc.add_table(rows=5, cols=2)
    ref_table.style = 'Table Grid'
    add_row(ref_table, "No. Rujukan Manual", f"HAS-MAN-001")
    add_row(ref_table, "Versi", "1.0")
    add_row(ref_table, "Tarikh Efektif", today)
    add_row(ref_table, "Disediakan oleh", s.halal_exec or "Eksekutif Halal")
    add_row(ref_table, "Diluluskan oleh", f"{s.approval_name} ({s.approval_title})" if s.approval_name else "Pihak Pengurusan Atasan")

    doc.add_page_break()

    # ── Senarai Kandungan ───────────────────────────────────────────────────
    add_section_header(doc, "SENARAI KANDUNGAN")
    toc_items = [
        "1.0  Profil Syarikat",
        "2.0  Objektif dan Skop",
        "3.0  Polisi Halal",
        "4.0  Jawatankuasa Halal Dalaman (JHD)",
        "5.0  Prosedur Operasi Standard (SOP)",
        "6.0  Rekod dan Dokumentasi",
        "7.0  Rujukan",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ── 1. Profil Syarikat ──────────────────────────────────────────────────
    add_section_header(doc, "1.0  PROFIL SYARIKAT")
    profile_table = doc.add_table(rows=6, cols=2)
    profile_table.style = 'Table Grid'
    add_row(profile_table, "Nama Syarikat", s.company_name or "[NAMA SYARIKAT]")
    add_row(profile_table, "Alamat", s.company_address or "[ALAMAT SYARIKAT]")
    add_row(profile_table, "No. Pendaftaran SSM", s.company_ssm or "[NO. SSM]")
    add_row(profile_table, "Saiz Syarikat", s.company_size or "[SAIZ SYARIKAT]")
    add_row(profile_table, "Industri", s.industry or "[INDUSTRI]")
    add_row(profile_table, "Sistem Halal", system)

    doc.add_paragraph()
    desc_para = doc.add_paragraph()
    desc_run = desc_para.add_run("Keterangan Produk / Perkhidmatan: ")
    desc_run.bold = True
    desc_run.font.size = Pt(11)
    doc.add_paragraph(s.product_desc or "[Sila huraikan produk/perkhidmatan syarikat]")

    doc.add_page_break()

    # ── 2. Objektif dan Skop ───────────────────────────────────────────────
    add_section_header(doc, "2.0  OBJEKTIF DAN SKOP")
    doc.add_paragraph(
        f"Manual {system} ini dibangunkan bertujuan untuk:"
    )
    objectives = [
        "Memastikan semua produk dan perkhidmatan yang dihasilkan mematuhi keperluan Pensijilan Halal Malaysia.",
        "Mewujudkan sistem kawalan halal yang berkesan dan berterusan di peringkat dalaman syarikat.",
        f"Memenuhi keperluan {system} seperti yang ditetapkan dalam MHMS 2020, MPPHM 2020, dan {ms}.",
        "Meningkatkan kepercayaan pengguna terhadap status halal produk syarikat.",
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    scope_header = doc.add_paragraph()
    scope_header.add_run("Skop:").bold = True
    doc.add_paragraph(
        f"Manual ini meliputi semua operasi {s.company_name or 'syarikat'} yang berkaitan dengan "
        f"pengeluaran, penyimpanan, pengendalian, dan pengedaran {s.product_desc or 'produk/perkhidmatan halal'}."
    )

    doc.add_page_break()

    # ── 3. Polisi Halal ─────────────────────────────────────────────────────
    add_section_header(doc, "3.0  POLISI HALAL")
    policy_text = s.halal_policy if s.halal_policy else (
        f"{s.company_name or 'Syarikat kami'} komited untuk memastikan semua produk dan perkhidmatan "
        f"yang dihasilkan adalah halal dan memenuhi keperluan Syariah serta standard yang ditetapkan "
        f"oleh JAKIM. Kami bertanggungjawab untuk melaksanakan, mengekalkan, dan menambah baik sistem "
        f"pengurusan halal secara berterusan bagi memenuhi keperluan pelanggan dan pihak berkepentingan."
    )
    policy_para = doc.add_paragraph(policy_text)
    policy_para.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    sign_para = doc.add_paragraph()
    sign_para.add_run(f"\n\n_________________________\n{s.approval_name or '[NAMA]'}\n{s.approval_title or 'Pengarah Urusan'}\n{today}")
    sign_para.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ── 4. Jawatankuasa Halal Dalaman ──────────────────────────────────────
    add_section_header(doc, "4.0  JAWATANKUASA HALAL DALAMAN (JHD)")
    doc.add_paragraph(
        "Jawatankuasa Halal Dalaman (JHD) dilantik secara rasmi oleh pengurusan atasan untuk "
        "membangun, melaksana, memantau, dan mengawal keberkesanan pelaksanaan HAS."
    )

    if s.ihc_members:
        doc.add_paragraph()
        jhd_table = doc.add_table(rows=1 + len(s.ihc_members), cols=3)
        jhd_table.style = 'Table Grid'
        header_row = jhd_table.rows[0]
        for i, h in enumerate(["Nama", "Jawatan dalam Syarikat", "Peranan dalam JHD"]):
            header_row.cells[i].text = h
            header_row.cells[i].paragraphs[0].runs[0].bold = True
        for i, member in enumerate(s.ihc_members, 1):
            jhd_table.rows[i].cells[0].text = member.get("name", "")
            jhd_table.rows[i].cells[1].text = member.get("position", "")
            jhd_table.rows[i].cells[2].text = member.get("role", "")
    else:
        placeholder_table = doc.add_table(rows=4, cols=3)
        placeholder_table.style = 'Table Grid'
        h_row = placeholder_table.rows[0]
        for i, h in enumerate(["Nama", "Jawatan dalam Syarikat", "Peranan dalam JHD"]):
            h_row.cells[i].text = h
            h_row.cells[i].paragraphs[0].runs[0].bold = True
        roles = [
            ("", "Pengarah Urusan", "Pengerusi JHD"),
            ("", "Pengurus Operasi", "Naib Pengerusi"),
            ("", "Eksekutif Halal", "Setiausaha"),
        ]
        for i, (n, pos, role) in enumerate(roles, 1):
            placeholder_table.rows[i].cells[0].text = "[NAMA]"
            placeholder_table.rows[i].cells[1].text = pos
            placeholder_table.rows[i].cells[2].text = role

    doc.add_page_break()

    # ── 5. SOP Sections ────────────────────────────────────────────────────
    add_section_header(doc, "5.0  PROSEDUR OPERASI STANDARD (SOP)")

    sops_to_include = s.selected_sops if s.selected_sops else [code for code, _ in MANDATORY_SOPS]

    sop_content = {
        "audit": {
            "title": "SOP Audit Halal Dalaman",
            "ref": "HAS-SOP-AHD-001",
            "objective": "Memastikan semua aktiviti berkaitan halal dilaksanakan mengikut keperluan yang ditetapkan melalui proses pengauditan dalaman yang sistematik.",
            "steps": [
                "1. Pelan Audit — Sediakan pelan audit tahunan oleh Eksekutif Halal.",
                "2. Notis Audit — Notifikasi kepada jabatan berkaitan sekurang-kurangnya 7 hari sebelum audit.",
                "3. Pelaksanaan — Jalankan audit menggunakan senarai semak yang diluluskan.",
                "4. Laporan — Dokumenkan penemuan dalam Laporan Audit Halal Dalaman.",
                "5. NCR — Keluarkan NCR untuk sebarang ketidakakuran yang dikenal pasti.",
                "6. Tindakan — Pantau pelaksanaan tindakan pembetulan dan pencegahan.",
                "7. Tutup — Sahkan penutupan NCR dan rekodkan dalam sistem.",
            ],
            "frequency": "Minimum sekali setahun. Audit luar jangka boleh dilakukan apabila perlu.",
        },
        "risk": {
            "title": "SOP Kawalan Risiko Halal (HCP)",
            "ref": "HAS-SOP-KRH-001",
            "objective": "Mengenal pasti, menganalisis, dan mengawal titik kawalan kritikal halal (Halal Control Point) dalam setiap proses pengeluaran.",
            "steps": [
                "1. Pembangunan Process Flow — Lakarkan aliran proses pengeluaran produk.",
                "2. Analisis Risiko — Kenal pasti ancaman kontaminasi halal di setiap tahap proses.",
                "3. Penetapan HCP — Tentukan mana titik kritikal yang memerlukan kawalan ketat.",
                "4. Had Kritikal — Tetapkan had boleh diterima bagi setiap HCP.",
                "5. Sistem Pemantauan — Bangunkan sistem pemantauan berterusan untuk setiap HCP.",
                "6. Tindakan Pembetulan — Dokumenkan prosedur tindakan apabila HCP melebihi had.",
                "7. Penyimpanan Rekod — Rekodkan semua aktiviti pemantauan HCP.",
            ],
            "frequency": "Pemantauan HCP dilakukan setiap hari semasa operasi.",
        },
        "rawmat": {
            "title": "SOP Kawalan Bahan Mentah",
            "ref": "HAS-SOP-KBM-001",
            "objective": "Memastikan semua bahan mentah yang digunakan dalam pengeluaran produk adalah halal dan berstatus yang ditetapkan.",
            "steps": [
                "1. Senarai Masterlist — Bangunkan dan selenggara Raw Material Masterlist.",
                "2. Penilaian Pembekal — Nilai pembekal berdasarkan kriteria halal yang ditetapkan.",
                "3. Penerimaan — Semak dokumen halal (sijil, CoA, CoO) semasa penerimaan bahan.",
                "4. Penyimpanan — Simpan bahan halal berasingan daripada bahan tidak halal / meragukan.",
                "5. Pelabelan — Label semua bahan mentah dengan status halal yang jelas.",
                "6. Pemantauan — Semak kelulusan/tamat tempoh sijil halal pembekal setiap suku tahun.",
                "7. Rekod — Selenggara rekod penerimaan dan kawalan bahan mentah.",
            ],
            "frequency": "Semakan harian semasa penerimaan. Semakan sijil halal setiap suku tahun.",
        },
        "training": {
            "title": "SOP Latihan Halal",
            "ref": "HAS-SOP-LH-001",
            "objective": "Memastikan semua pekerja yang terlibat dalam operasi mempunyai kesedaran dan kefahaman yang mencukupi tentang keperluan halal.",
            "steps": [
                "1. Keperluan Latihan — Kenal pasti keperluan latihan halal untuk setiap jawatan.",
                "2. Pelan Latihan — Bangunkan pelan latihan halal tahunan.",
                "3. Latihan Kesedaran — Jalankan latihan kesedaran halal untuk pekerja baru dalam 3 bulan pertama.",
                "4. Latihan Berkala — Latihan halal kepada semua pekerja sekurang-kurangnya sekali dalam 3 tahun.",
                "5. Latihan JHD — Ahli JHD mendapat latihan kompetensi daripada Penyedia Latihan Halal (PLH) berdaftar HPB.",
                "6. Penilaian — Jalankan penilaian untuk mengukur keberkesanan latihan.",
                "7. Rekod — Rekodkan kehadiran, penilaian, dan sijil latihan.",
            ],
            "frequency": "Latihan kesedaran: pekerja baru dalam 3 bulan. Latihan berkala: sekali dalam 3 tahun.",
        },
        "trace": {
            "title": "SOP Kebolehkesanan (Traceability)",
            "ref": "HAS-SOP-KBS-001",
            "objective": "Memastikan setiap produk boleh dikesan ke belakang (bahan mentah) dan ke hadapan (pengedaran) pada bila-bila masa.",
            "steps": [
                "1. Sistem Penomboran — Wujudkan sistem nombor lot/batch yang unik untuk setiap pengeluaran.",
                "2. Rekod Pengeluaran — Rekodkan nombor lot bahan mentah yang digunakan dalam setiap batch.",
                "3. Rekod Pengedaran — Rekodkan nombor lot produk siap yang dihantar kepada setiap pelanggan.",
                "4. Ujian Kebolehkesanan — Jalankan ujian kebolehkesanan sekurang-kurangnya sekali setahun.",
                "5. Prosedur Penarikan — Bangunkan prosedur penarikan produk (product recall) yang jelas.",
                "6. Masa Tindak Balas — Pastikan keseluruhan lot boleh dikesan dalam masa 4 jam.",
                "7. Rekod — Simpan semua rekod kebolehkesanan sekurang-kurangnya 3 tahun.",
            ],
            "frequency": "Rekod harian. Ujian kebolehkesanan sekali setahun.",
        },
        "review": {
            "title": "SOP Semakan HAS",
            "ref": "HAS-SOP-SH-001",
            "objective": "Memastikan sistem HAS sentiasa relevan, mencukupi, dan berkesan melalui semakan berkala oleh pihak pengurusan.",
            "steps": [
                "1. Jadual Semakan — JHD menjadualkan semakan HAS sekurang-kurangnya sekali setahun.",
                "2. Input Semakan — Kumpulkan input: hasil audit, NCR, aduan, perubahan proses, perubahan standard.",
                "3. Penilaian — Nilai keberkesanan HAS berdasarkan input yang dikumpulkan.",
                "4. Keputusan — Tentukan tindakan penambahbaikan yang diperlukan.",
                "5. Kemaskini Dokumen — Kemaskini manual dan SOP jika ada perubahan.",
                "6. Rekod — Dokumenkan keputusan semakan dalam minit mesyuarat.",
            ],
            "frequency": "Minimum sekali setahun atau apabila berlaku perubahan signifikan.",
        },
        "lab": {
            "title": "SOP Analisis Makmal",
            "ref": "HAS-SOP-AM-001",
            "objective": "Menetapkan prosedur untuk pengambilan sampel dan analisis makmal bagi tujuan pengesahan status halal.",
            "steps": [
                "1. Pengenalpastian Keperluan — Kenal pasti produk/bahan yang memerlukan analisis makmal.",
                "2. Pemilihan Makmal — Gunakan makmal yang diiktiraf oleh JAKIM / akreditasi SAMM.",
                "3. Pengambilan Sampel — Ambil sampel mengikut prosedur yang ditetapkan.",
                "4. Penghantaran — Hantar sampel ke makmal dalam tempoh yang sesuai.",
                "5. Keputusan — Terima dan simpan laporan analisis makmal.",
                "6. Tindakan — Ambil tindakan segera jika keputusan menunjukkan kontaminasi.",
                "7. Rekod — Fail semua laporan analisis makmal sekurang-kurangnya 3 tahun.",
            ],
            "frequency": "Bergantung kepada keperluan: produk berisiko tinggi — setiap suku tahun. Produk lain — setahun sekali.",
        },
        "sertu": {
            "title": "SOP Sertu",
            "ref": "HAS-SOP-STU-001",
            "objective": "Menetapkan prosedur pelaksanaan sertu (penyucian daripada najis babi/anjing) mengikut hukum syarak.",
            "steps": [
                "1. Pengenalpastian — Kenal pasti peralatan/permukaan yang memerlukan sertu.",
                "2. Permohonan — Dapatkan kelulusan daripada Eksekutif Halal sebelum sertu.",
                "3. Pelaksanaan — Laksanakan sertu mengikut prosedur: bersihkan kotoran, basuh 7 kali dengan air, sekali dengan air tanah.",
                "4. Pengesahan — Eksekutif Halal mengesahkan pelaksanaan sertu.",
                "5. Rekod — Rekodkan tarikh, sebab, dan petugas yang melaksanakan sertu.",
                "6. Laporan — Laporkan kepada JHD dalam mesyuarat bulanan.",
            ],
            "frequency": "Dilaksanakan apabila berlaku kontaminasi najis babi/anjing.",
        },
    }

    for i, (code, sop_name) in enumerate(MANDATORY_SOPS, 1):
        if code in sops_to_include:
            content = sop_content.get(code, {})
            add_section_header(doc, f"5.{i}  {content.get('title', sop_name)}", bg_color=(15, 110, 86))

            ref_doc_table = doc.add_table(rows=2, cols=2)
            ref_doc_table.style = 'Table Grid'
            add_row(ref_doc_table, "No. Rujukan", content.get("ref", f"HAS-SOP-{code.upper()}-001"))
            add_row(ref_doc_table, "Frekuensi", content.get("frequency", "Seperti yang ditetapkan"))

            doc.add_paragraph()
            obj_p = doc.add_paragraph()
            obj_p.add_run("Objektif: ").bold = True
            obj_p.add_run(content.get("objective", ""))

            doc.add_paragraph()
            proc_p = doc.add_paragraph()
            proc_p.add_run("Prosedur:").bold = True

            for step in content.get("steps", []):
                p = doc.add_paragraph(step, style='List Bullet')
                p.runs[0].font.size = Pt(11)

            doc.add_paragraph()

    doc.add_page_break()

    # ── 6. Rekod ────────────────────────────────────────────────────────────
    add_section_header(doc, "6.0  REKOD DAN DOKUMENTASI")
    doc.add_paragraph(
        "Semua rekod berkaitan HAS hendaklah disimpan, dikemaskini, dan mudah dirujuk. "
        "Tempoh penyimpanan: minimum 3 tahun untuk permohonan baru; minimum 3 bulan untuk "
        "permohonan pembaharuan."
    )

    doc.add_paragraph()
    records_table = doc.add_table(rows=1, cols=4)
    records_table.style = 'Table Grid'
    h_row = records_table.rows[0]
    for i, h in enumerate(["Nama Rekod", "Berkaitan SOP", "Frekuensi", "Tempoh Simpan"]):
        h_row.cells[i].text = h
        h_row.cells[i].paragraphs[0].runs[0].bold = True
        h_row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    freq_map = {"audit": "Bulanan", "risk": "Harian", "rawmat": "Harian", "training": "Tahunan",
                "trace": "Harian", "review": "Tahunan", "lab": "Suku Tahun", "sertu": "Atas Keperluan"}

    for code in sops_to_include:
        records_list = RECORDS.get(code, [])
        sop_name = dict(MANDATORY_SOPS).get(code, code)
        for rec in records_list:
            row = records_table.add_row()
            row.cells[0].text = rec
            row.cells[1].text = sop_name
            row.cells[2].text = freq_map.get(code, "Seperti ditetapkan")
            row.cells[3].text = "3 tahun"
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── 7. Rujukan ──────────────────────────────────────────────────────────
    add_section_header(doc, "7.0  RUJUKAN")
    references = [
        f"1. Sistem Pengurusan Halal Malaysia (MHMS) 2020 — JAKIM",
        f"2. Manual Prosedur Pensijilan Halal Malaysia (MPPHM) 2020 — JAKIM",
        f"3. {ms} — Jabatan Standard Malaysia",
        "4. Fatwa-fatwa berkaitan halal yang dikeluarkan oleh Majlis Fatwa Kebangsaan",
        "5. Undang-undang dan peraturan berkaitan yang diguna pakai di Malaysia",
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    # ── Save to buffer ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─── UI Steps ─────────────────────────────────────────────────────────────────

def render_step1():
    header("Profil Syarikat", 1)
    
    st.session_state.company_name = st.text_input(
        "Nama syarikat (seperti dalam SSM)",
        value=st.session_state.company_name,
        placeholder="cth: Syarikat ABC Sdn. Bhd."
    )
    st.session_state.company_address = st.text_area(
        "Alamat syarikat",
        value=st.session_state.company_address,
        placeholder="No. 1, Jalan Contoh, 50000 Kuala Lumpur",
        height=80
    )
    st.session_state.company_ssm = st.text_input(
        "No. pendaftaran SSM",
        value=st.session_state.company_ssm,
        placeholder="cth: 202001012345 (1234567-A)"
    )

    st.markdown("---")
    st.session_state.company_size = st.selectbox(
        "Saiz syarikat",
        options=[""] + list(COMPANY_SIZES.keys()),
        index=0 if not st.session_state.company_size else list(COMPANY_SIZES.keys()).index(st.session_state.company_size) + 1,
        help="Saiz syarikat menentukan sama ada perlu IHCS atau HAS penuh."
    )

    if st.session_state.company_size:
        size_code = COMPANY_SIZES[st.session_state.company_size]
        system = "IHCS (Sistem Kawalan Halal Dalaman)" if size_code == "micro" else "HAS Penuh (Sistem Jaminan Halal)"
        css_class = "info-box" if size_code == "micro" else "warning-box"
        st.markdown(f'<div class="{css_class}">🏷 Saiz ini memerlukan: <strong>{system}</strong></div>', unsafe_allow_html=True)

    st.markdown("---")
    col1, col2 = st.columns([1, 5])
    with col2:
        if st.button("Seterusnya →", type="primary", use_container_width=False):
            if not st.session_state.company_name:
                st.error("Sila masukkan nama syarikat.")
            elif not st.session_state.company_size:
                st.error("Sila pilih saiz syarikat.")
            else:
                st.session_state.step = 2
                st.rerun()

def render_step2():
    header("Industri & Produk", 2)

    st.session_state.industry = st.selectbox(
        "Jenis industri",
        options=[""] + list(INDUSTRIES.keys()),
        index=0 if not st.session_state.industry else list(INDUSTRIES.keys()).index(st.session_state.industry) + 1,
        help="Industri menentukan Malaysian Standard (MS) yang dipakai dalam manual."
    )

    if st.session_state.industry:
        ms = INDUSTRIES[st.session_state.industry]["ms"]
        st.markdown(f'<div class="info-box">📋 Malaysian Standard yang terpakai: <strong>{ms}</strong></div>', unsafe_allow_html=True)

    st.session_state.product_desc = st.text_area(
        "Huraikan produk / perkhidmatan syarikat",
        value=st.session_state.product_desc,
        placeholder="cth: Pengeluaran kuih-muih tradisional seperti kuih lapis, onde-onde, dan karipap untuk pasaran runcit dan borong.",
        height=100
    )

    st.markdown("---")
    col_back, col_next = st.columns([1, 1])
    with col_back:
        if st.button("← Balik", use_container_width=True):
            st.session_state.step = 1
            st.rerun()
    with col_next:
        if st.button("Seterusnya →", type="primary", use_container_width=True):
            if not st.session_state.industry:
                st.error("Sila pilih industri.")
            else:
                st.session_state.step = 3
                st.rerun()

def render_step3():
    header("Jawatankuasa Halal Dalaman (JHD)", 3)

    st.markdown('<div class="info-box">JHD mesti ditubuhkan secara rasmi dan bertulis oleh pengurusan atasan syarikat.</div>', unsafe_allow_html=True)

    st.session_state.halal_exec = st.text_input(
        "Nama Eksekutif Halal (Setiausaha JHD)",
        value=st.session_state.halal_exec,
        placeholder="Nama penuh Eksekutif Halal yang berdaftar HPB"
    )

    st.subheader("Ahli Jawatankuasa")
    st.caption("Tambah semua ahli JHD. Minimum: Eksekutif Halal + wakil pengurusan atasan.")

    if "ihc_members" not in st.session_state or not st.session_state.ihc_members:
        st.session_state.ihc_members = [
            {"name": "", "position": "Pengarah Urusan", "role": "Pengerusi JHD"},
            {"name": "", "position": "Eksekutif Halal", "role": "Setiausaha JHD"},
        ]

    updated_members = []
    for i, member in enumerate(st.session_state.ihc_members):
        with st.expander(f"Ahli {i+1}: {member['position']}", expanded=True):
            cols = st.columns([2, 2, 2])
            name = cols[0].text_input("Nama", value=member["name"], key=f"mem_name_{i}", placeholder="Nama penuh")
            position = cols[1].text_input("Jawatan dalam syarikat", value=member["position"], key=f"mem_pos_{i}")
            role = cols[2].text_input("Peranan dalam JHD", value=member["role"], key=f"mem_role_{i}")
            updated_members.append({"name": name, "position": position, "role": role})

    st.session_state.ihc_members = updated_members

    col_add, _ = st.columns([1, 3])
    with col_add:
        if st.button("+ Tambah ahli"):
            st.session_state.ihc_members.append({"name": "", "position": "", "role": "Ahli JHD"})
            st.rerun()

    st.markdown("---")
    col_back, col_next = st.columns([1, 1])
    with col_back:
        if st.button("← Balik", use_container_width=True):
            st.session_state.step = 2
            st.rerun()
    with col_next:
        if st.button("Seterusnya →", type="primary", use_container_width=True):
            st.session_state.step = 4
            st.rerun()

def render_step4():
    header("SOP, HCP & Polisi Halal", 4)
    system = get_system_type()

    st.subheader("Pilih SOP yang perlu dibangunkan")
    if system == "IHCS":
        st.markdown('<div class="info-box">Sebagai syarikat Mikro/Kecil (IHCS), SOP minimum adalah: Polisi Halal, Kawalan Bahan Mentah, dan Kebolehkesanan.</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="warning-box">Sebagai syarikat HAS penuh, semua 8 SOP di bawah adalah wajib mengikut MHMS 2020.</div>', unsafe_allow_html=True)

    sop_selections = []
    cols = st.columns(2)
    for i, (code, label) in enumerate(MANDATORY_SOPS):
        default_checked = system == "HAS" or code in ["rawmat", "trace", "risk"]
        checked = cols[i % 2].checkbox(label, value=code in st.session_state.selected_sops or default_checked, key=f"sop_{code}")
        if checked:
            sop_selections.append(code)
    st.session_state.selected_sops = sop_selections

    st.markdown("---")
    st.subheader("Polisi Halal")
    st.caption("Tulis polisi halal syarikat atau biarkan kosong untuk guna templat.")
    st.session_state.halal_policy = st.text_area(
        "Polisi Halal (opsional — templat akan dijana jika kosong)",
        value=st.session_state.halal_policy,
        placeholder="cth: Kami di [nama syarikat] komited untuk memastikan...",
        height=120,
        label_visibility="collapsed"
    )

    st.markdown("---")
    st.subheader("Maklumat Pengesahan Dokumen")
    col_a, col_b = st.columns(2)
    st.session_state.approval_name = col_a.text_input("Nama pelulus (Pengurusan Atasan)", value=st.session_state.approval_name)
    st.session_state.approval_title = col_b.text_input("Jawatan", value=st.session_state.approval_title, placeholder="cth: Pengarah Urusan")
    st.session_state.doc_controller = st.text_input("Pengawal dokumen", value=st.session_state.doc_controller, placeholder="cth: Eksekutif Halal")

    st.markdown("---")
    col_back, col_next = st.columns([1, 1])
    with col_back:
        if st.button("← Balik", use_container_width=True):
            st.session_state.step = 3
            st.rerun()
    with col_next:
        if st.button("Semak & Jana Manual →", type="primary", use_container_width=True):
            st.session_state.step = 5
            st.rerun()

def render_step5():
    header("Semak & Jana Manual HAS", 5)
    s = st.session_state
    system = get_system_type()
    ms = get_ms_standard()

    st.subheader("Ringkasan")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Syarikat", s.company_name or "—")
    col2.metric("Sistem", system)
    col3.metric("Industri", s.industry.split("/")[0].strip() if s.industry else "—")
    col4.metric("MS Standard", ms)

    st.markdown("---")
    st.subheader("Kandungan manual yang akan dijana")

    base_sections = ["Profil Syarikat", "Objektif & Skop", "Polisi Halal", "Jawatankuasa Halal Dalaman"]
    all_items = base_sections + [dict(MANDATORY_SOPS).get(code, code) for code in s.selected_sops]
    all_items += ["Rekod & Dokumentasi", "Rujukan (MHMS, MPPHM, MS Standard)"]

    for item in all_items:
        st.markdown(f"✅ {item}")

    st.markdown("---")
    st.subheader("Rujukan silang wajib")
    st.markdown(f"""
    - 📄 MHMS 2020 — Malaysian Halal Management System
    - 📄 MPPHM 2020 — Manual Prosedur Pensijilan Halal Malaysia  
    - 📄 {ms} — Malaysian Standard (ikut industri)
    """)

    st.markdown("---")
    col_back, col_gen = st.columns([1, 2])
    with col_back:
        if st.button("← Balik & Edit", use_container_width=True):
            st.session_state.step = 4
            st.rerun()
    with col_gen:
        if st.button("⬇ Jana & Muat Turun Manual (Word)", type="primary", use_container_width=True):
            with st.spinner("Menjana manual HAS..."):
                buf = generate_has_manual()
            fname = f"Manual_HAS_{s.company_name.replace(' ', '_') or 'Syarikat'}_{datetime.now().strftime('%Y%m%d')}.docx"
            st.download_button(
                label="📥 Klik untuk muat turun Manual HAS (.docx)",
                data=buf,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.success(f"Manual HAS berjaya dijana! Fail: {fname}")
            st.info("💡 Manual ini adalah draf permulaan. Sila semak bersama Halal Advisor / Eksekutif Halal sebelum diserahkan kepada JAKIM.")

# ─── Main Layout ───────────────────────────────────────────────────────────────

st.markdown('<h1 style="font-family: DM Serif Display, serif; font-size: 28px; color: #0F6E56; margin-bottom: 4px;">HAS Manual Generator</h1>', unsafe_allow_html=True)
st.markdown('<p style="color: #888780; font-size: 14px; margin-bottom: 1.5rem;">Platform Pembangunan Manual Halal Assurance System (MHMS 2020 / JAKIM)</p>', unsafe_allow_html=True)

progress_bar()

step = st.session_state.step
if step == 1:
    render_step1()
elif step == 2:
    render_step2()
elif step == 3:
    render_step3()
elif step == 4:
    render_step4()
elif step == 5:
    render_step5()
